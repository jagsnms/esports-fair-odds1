"""
Convert GRID_telemetry.docx to docs/GRID_telemetry_reference.md.
Preserves structure; redacts likely secrets.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# Namespace for Word
W_P = qn("w:p")
W_TBL = qn("w:tbl")
W_TC = qn("w:tc")


def cell_text(cell) -> str:
    return " ".join(p.text or "" for p in cell.paragraphs).strip()


def table_to_md(table) -> str:
    rows = []
    for tr in table.rows:
        cells = [cell_text(cell) for cell in tr.cells]
        rows.append("| " + " | ".join(c.replace("|", "\\|") for c in cells) + " |")
    if not rows:
        return ""
    # First row as header
    out = [rows[0], "| " + " | ".join("---" for _ in table.columns) + " |"]
    out.extend(rows[1:])
    return "\n".join(out)


def redact(text: str) -> str:
    if not text:
        return text
    # Redact common secret patterns (case-insensitive)
    patterns = [
        (r'\b(?:api[_-]?key|apikey)\s*[:=]\s*["\']?[^\s"\']+["\']?', r'\1 = [REDACTED]'),
        (r'\b(?:token|secret|password|auth)\s*[:=]\s*["\']?[^\s"\']+["\']?', r'\1 = [REDACTED]'),
        (r'["\']?(?:sk|pk)_[a-zA-Z0-9]{20,}["\']?', '[REDACTED]'),
        (r'\b[A-Za-z0-9]{32,}\b', lambda m: m.group(0) if not re.match(r'^[a-fA-F0-9]{32,}$', m.group(0)) else '[REDACTED]'),
    ]
    # Simpler: replace any long alphanumeric that looks like key/token
    text = re.sub(r'\b(?:api[_-]?key|token|secret)\s*[:=]\s*\S+', '[REDACTED]', text, flags=re.I)
    return text


def main() -> None:
    repo = Path(__file__).resolve().parent.parent
    docx_path = repo / "GRID_telemetry.docx"
    out_path = repo / "docs" / "GRID_telemetry_reference.md"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(docx_path))
    body = doc.element.body
    para_idx = 0
    tbl_idx = 0
    lines = [
        "Source: GRID_telemetry.docx (converted for searchability). Do not store secrets here.",
        "",
        "---",
        "",
    ]
    last_was_table = False

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            if para_idx >= len(doc.paragraphs):
                continue
            p = doc.paragraphs[para_idx]
            para_idx += 1
            text = (p.text or "").strip().replace("\xa0", " ")
            text = redact(text)
            if not text:
                lines.append("")
                continue
            # Short line after table or at start -> treat as heading
            if (last_was_table or (len(lines) <= 6 and not last_was_table)) and len(text) < 80 and "\n" not in text:
                if text.lower() not in ("objects", "queries", "mutations") and not text.endswith("."):
                    lines.append("## " + text)
                    lines.append("")
                    last_was_table = False
                    continue
            lines.append(text)
            lines.append("")
            last_was_table = False
        elif tag == "tbl":
            if tbl_idx < len(doc.tables):
                tbl = doc.tables[tbl_idx]
                tbl_idx += 1
                md = table_to_md(tbl)
                if md:
                    lines.append(md)
                    lines.append("")
                last_was_table = True
            else:
                last_was_table = False

    out_path.write_text("\n".join(lines), encoding="utf-8")
    print("Wrote", out_path)


if __name__ == "__main__":
    main()
